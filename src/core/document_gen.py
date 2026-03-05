"""文档导出模块 — Markdown / Word / Excel 导出"""

from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt

from .config import settings


class DocumentGenerator:
    """多格式文档生成器"""

    def __init__(self, output_dir: Path | None = None):
        self._output_dir = output_dir or settings.output_dir

    def save_markdown(self, content: str, filename: str) -> Path:
        """保存 Markdown 文件"""
        path = self._output_dir / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return path

    def save_word(self, content: str, filename: str, title: str = "") -> Path:
        """将 Markdown 内容导出为 Word 文档（简单转换）"""
        doc = Document()
        if title:
            doc.add_heading(title, level=0)

        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                # 表格行直接作为普通段落（简单处理）
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(10)
            elif line.strip():
                doc.add_paragraph(line)

        path = self._output_dir / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    def save_excel(self, data: list[dict[str, Any]], filename: str, sheet_name: str = "Sheet1") -> Path:
        """保存数据为 Excel 文件"""
        df = pd.DataFrame(data)
        path = self._output_dir / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(str(path), sheet_name=sheet_name, index=False)
        return path

    def save_multi_sheet_excel(self, sheets: dict[str, list[dict]], filename: str) -> Path:
        """保存多 Sheet 的 Excel 文件"""
        path = self._output_dir / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        with pd.ExcelWriter(str(path)) as writer:
            for sheet_name, data in sheets.items():
                df = pd.DataFrame(data)
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        return path
